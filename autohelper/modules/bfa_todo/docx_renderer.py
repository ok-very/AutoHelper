"""
Render the BFA To Do List as a .docx via Word COM + python-docx post-processing.

Pipeline:
  1. Pre-process index_pasteable.html (reorder sections, update date, strip comments)
  2. Word COM opens the HTML and saves as .docx (handles inline CSS natively)
  3. Post-process with python-docx: insert section breaks at text anchors to match
     the original document's 6-section structure (1-col heading, 2-col policies,
     2-col proposals, transition sections, 1-col projects), fix margins to ~0.8in
  4. Upload to Google Drive via files.update
"""

from __future__ import annotations

import logging
import re
import tempfile
from datetime import datetime
from pathlib import Path

from autohelper.shared.paths import data_dir

logger = logging.getLogger(__name__)

# Margins matching the original: ~0.8 inches = 1134 twips
_MARGIN_TWIPS = "1134"


# ---------------------------------------------------------------------------
# Stage 1: Pre-process HTML
# ---------------------------------------------------------------------------

def _get_on_hold_uids() -> set[str]:
    try:
        from .pipeline.repo import BfaProjectRepo
        repo = BfaProjectRepo()
        projects = repo.get_all(include_archived=False)
        return {p["uid"] for p in projects if p.get("status") == "on_hold"}
    except Exception:
        logger.warning("Could not load project statuses from DB")
        return set()


def _prepare_html_for_word(html: str) -> str:
    """Reorder sections, update heading date, strip comments."""
    from bs4 import BeautifulSoup, Comment

    soup = BeautifulSoup(html, "lxml")
    body = soup.body
    if not body:
        return html

    on_hold_uids = _get_on_hold_uids()

    # -- Extract sections by type --
    preambles = []
    preamble_lists = []
    project_sections = []

    for section in body.find_all("section", recursive=False):
        classes = section.get("class", [])
        if "bfa-preamble-lists" in classes:
            preamble_lists.append(section.extract())
        elif "bfa-preamble" in classes:
            preambles.append(section.extract())
        elif "bfa-project" in classes:
            project_sections.append(section.extract())

    # -- Separate active vs on-hold --
    active = []
    on_hold = []
    for sec in project_sections:
        uid = sec.get("data-project-uid", "")
        if uid in on_hold_uids:
            on_hold.append(sec)
        else:
            active.append(sec)

    # -- Rebuild body in correct order --
    body.clear()

    for sec in preambles:
        body.append(sec)
    for sec in preamble_lists:
        if sec.get("style"):
            sec["style"] = re.sub(
                r"column-count:\s*\d+;?\s*|column-gap:[^;]+;?\s*|column-rule:[^;]+;?\s*",
                "", sec["style"],
            )
        body.append(sec)
    for sec in active:
        body.append(sec)

    if on_hold:
        sep = soup.new_tag("p")
        sep["style"] = (
            "font-weight:bold; text-align:center; font-size:11pt; "
            "font-family:Calibri; border-top:2px solid #000; "
            "padding-top:6pt; margin-top:12pt;"
        )
        sep.string = f"{'─' * 40}  ON HOLD  {'─' * 40}"
        body.append(sep)
        for sec in on_hold:
            body.append(sec)

    # -- Extract project UID + header hints for SDT embedding --
    uid_map: list[dict[str, str]] = []
    for sec in active:
        h3 = sec.find("h3")
        uid_map.append({
            "uid": sec.get("data-project-uid", ""),
            "header": h3.get_text(strip=True) if h3 else "",
        })
    for sec in on_hold:
        h3 = sec.find("h3")
        uid_map.append({
            "uid": sec.get("data-project-uid", ""),
            "header": h3.get_text(strip=True) if h3 else "",
        })

    # -- Update heading date --
    now = datetime.now()
    today_str = f"{now.strftime('%B')} {now.day}, {now.year}"
    for preamble in body.find_all("section", class_="bfa-preamble"):
        for span in preamble.find_all("span"):
            if span.string and "To Do List" in span.string:
                span.string = re.sub(
                    r"(To Do List\s+)\w+ \d{1,2},?\s*\d{4}",
                    rf"\g<1>{today_str}",
                    span.string,
                )
                break
        else:
            for span in preamble.find_all("span"):
                if span.string and "To Do List" in span.string:
                    nxt = span.find_next_sibling("span")
                    if nxt and nxt.string and re.match(r"\w+ \d{1,2},?\s*\d{4}", nxt.string.strip()):
                        nxt.string = today_str
                    break

    # -- Flatten tables inside project sections only --
    # Preamble tables (PAC schedule) stay as Word tables.
    for table in body.find_all("table"):
        parent_section = table.find_parent("section")
        if parent_section and "bfa-project" in parent_section.get("class", []):
            replacement = []
            for tr in table.find_all("tr"):
                cells = tr.find_all(["td", "th"])
                row_text = "\t".join(c.get_text(strip=True) for c in cells)
                if row_text.strip():
                    p_tag = soup.new_tag("p")
                    p_tag["style"] = "font-size:11pt; font-family:Calibri; margin:0; padding:0;"
                    p_tag.string = row_text
                    replacement.append(p_tag)
            if replacement:
                for p_tag in replacement:
                    table.insert_before(p_tag)
            table.decompose()

    # -- Fix font sizes: 10pt → 11pt for PROJECT text only --
    # The base.css forces 10pt for compact web view, but the original .docx is
    # 11pt throughout project sections.  Preamble text keeps its original sizes
    # (8pt for policy details, inherited for headings).
    for section in body.find_all("section"):
        if "bfa-project" not in section.get("class", []):
            continue
        for el in section.find_all(["p", "li", "span"]):
            style = el.get("style", "")
            if "font-size:10pt" in style or "font-size: 10pt" in style:
                el["style"] = re.sub(r"font-size:\s*10pt", "font-size:11pt", style)

    # -- Strip CSS class attributes --
    # Prevents Word COM from creating c0/c18/c21 paragraph styles.
    for el in body.find_all(True):
        if el.get("class"):
            del el["class"]

    # -- Strip HTML comments --
    for comment in soup.find_all(string=lambda t: isinstance(t, Comment)):
        comment.extract()

    return str(soup), uid_map


# ---------------------------------------------------------------------------
# Stage 2: Word COM conversion
# ---------------------------------------------------------------------------

def _html_to_docx_via_word(html_path: Path, docx_path: Path) -> Path:
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

        doc = word.Documents.Open(
            str(html_path.resolve()),
            ConfirmConversions=False,
            ReadOnly=False,
            AddToRecentFiles=False,
        )
        doc.SaveAs2(str(docx_path.resolve()), FileFormat=16)
        logger.info("Word converted HTML -> DOCX: %s", docx_path.name)
    finally:
        if doc:
            try:
                doc.Close(SaveChanges=False)
            except Exception:
                pass
        if word:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()

    return docx_path


def _docx_to_pdf(docx_path: Path, pdf_path: Path) -> Path:
    """Convert .docx to PDF via Word COM."""
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.SaveAs2(str(pdf_path.resolve()), FileFormat=17)  # wdFormatPDF
    finally:
        if doc:
            try:
                doc.Close(SaveChanges=False)
            except Exception:
                pass
        if word:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()

    return pdf_path


# ---------------------------------------------------------------------------
# Stage 3: Post-process
#
# Structured as three clean phases:
#   Phase A — Read & analyze: find text anchors, no mutations
#   Phase B — XML mutations: section breaks, blank paragraph insertion
#   Phase C — Save, reopen, apply python-docx-level formatting
#             (styles, spacing, page breaks) on a clean document
# ---------------------------------------------------------------------------

def _postprocess_docx(docx_path: Path, uid_map: list[dict[str, str]] | None = None) -> None:
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree
    from copy import deepcopy

    # ── Step 1: Find all anchors on clean paragraph list ────────────────

    doc = Document(str(docx_path))
    paras = doc.paragraphs

    idx_city_policies = None
    idx_proposals = None
    idx_first_project = None

    for i, p in enumerate(paras):
        t = p.text.strip()
        if idx_city_policies is None and "Vancouver:" in t and "Checklist" in t:
            idx_city_policies = i
        if idx_proposals is None and t.startswith("Proposals") and "New" in t:
            idx_proposals = i
        if idx_first_project is None and t.startswith("(") and ":" in t:
            if any(kw in t for kw in ["Artwork", "Total", "Install"]):
                idx_first_project = i

    if idx_city_policies is None:
        logger.warning("Anchor missing: city policies start")
        return
    if idx_first_project is None:
        logger.warning("Anchor missing: first project header")
        return

    project_header_indices = []
    for i in range(idx_first_project, len(paras)):
        t = paras[i].text.strip()
        if t.startswith("(") and ":" in t and any(k in t for k in ["Artwork", "Total", "Install"]):
            project_header_indices.append(i)

    logger.info(
        "Step 1 — anchors: city_policies=%d, proposals=%s, first_project=%d, project_blocks=%d",
        idx_city_policies, idx_proposals, idx_first_project, len(project_header_indices),
    )

    # ── Step 2: Normalize styles (paragraphs are still proper objects) ───

    normal_style = doc.styles["Normal"]
    heading_styles = {"Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4"}
    style_changes = 0
    for p in paras:
        if p.style and p.style.name not in heading_styles:
            p.style = normal_style
            style_changes += 1

    # Zero out CSS-imported spacing on preamble paragraphs.
    # Original preamble has sa=None/0, sb=None/0, ls=1.0.
    # Word COM imports HTML margins as space_before/space_after (e.g. 63500 EMU)
    # which inflates the 2-col sections and causes column spill.
    #
    # Also fix bullet list indentation: Word COM creates wide indents (637540)
    # from HTML <ul><li>.  Original uses tight hanging indent:
    #   left_indent=179705, first_line_indent=-228600
    list_fixes = 0
    for i in range(idx_first_project):
        p = paras[i]
        pf = p.paragraph_format
        if pf.space_before and pf.space_before > 0:
            pf.space_before = 0
        if pf.space_after and pf.space_after > 0:
            pf.space_after = 0

        # Kill autospacing (Word adds automatic space around HTML block elements)
        pPr_el = p._element.find(qn("w:pPr"))
        if pPr_el is not None:
            spacing = pPr_el.find(qn("w:spacing"))
            if spacing is not None:
                spacing.set(qn("w:beforeAutospacing"), "0")
                spacing.set(qn("w:afterAutospacing"), "0")

        # Fix list paragraph indentation
        if pPr_el is not None and pPr_el.find(qn("w:numPr")) is not None:
            pf.left_indent = 179705
            pf.first_line_indent = -228600
            list_fixes += 1

    # Page break before "Proposals (New Projects)" so it gets its own page
    if idx_proposals is not None:
        paras[idx_proposals].paragraph_format.page_break_before = True

    # Insert blank paragraph separators between preamble list sub-blocks.
    # Original has blanks before: Art Plans, MASTER PUBLIC ART PLAN,
    # Longlists, Artist Contracts (To be drafted), Artist Contracts (In Negotiation)
    _PREAMBLE_BLOCK_ANCHORS = [
        "Art Plans to be drafted",
        "MASTER PUBLIC ART PLAN",
        "Longlists",
        "Proposals Pending",
        "Artist Contracts (To be drafted)",
        "Artist Contracts (In Negotiation)",
        "Final Documents and Installation",
    ]
    blank_xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    blanks_inserted = 0
    for i in range(idx_first_project):
        t = paras[i].text.strip()
        for anchor in _PREAMBLE_BLOCK_ANCHORS:
            if t.startswith(anchor):
                paras[i]._element.addprevious(
                    etree.fromstring(blank_xml)
                )
                blanks_inserted += 1
                break

    logger.info("Step 2 — styles: normalized %d, zeroed preamble spacing, fixed %d list indents, %d preamble blanks, proposals page break",
                style_changes, list_fixes, blanks_inserted)

    # Fix bullet character: change abstract numbering definitions used by
    # preamble lists from • (Symbol \uf0b7) to ● (BLACK CIRCLE U+25CF)
    # matching the original document.
    numbering_part = doc.part.numbering_part
    if numbering_part:
        # Collect abstractNumIds referenced by preamble list paragraphs
        preamble_abs_ids = set()
        for i in range(idx_first_project):
            pPr = paras[i]._element.find(qn("w:pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                if numPr is not None:
                    numId_el = numPr.find(qn("w:numId"))
                    if numId_el is not None:
                        nid = numId_el.get(qn("w:val"))
                        # Look up the abstractNumId for this numId
                        for num in numbering_part._element.findall(qn("w:num")):
                            if num.get(qn("w:numId")) == nid:
                                absRef = num.find(qn("w:abstractNumId"))
                                if absRef is not None:
                                    preamble_abs_ids.add(absRef.get(qn("w:val")))
                                break

        # Patch those abstract definitions: lvl 0 bullet text → ●
        for absNum in numbering_part._element.findall(qn("w:abstractNum")):
            if absNum.get(qn("w:abstractNumId")) in preamble_abs_ids:
                for lvl in absNum.findall(qn("w:lvl")):
                    if lvl.get(qn("w:ilvl")) == "0":
                        lvlText = lvl.find(qn("w:lvlText"))
                        if lvlText is not None:
                            lvlText.set(qn("w:val"), "●")
                        # Remove font override so it inherits (original has no font spec)
                        rPr = lvl.find(qn("w:rPr"))
                        if rPr is not None:
                            rFonts = rPr.find(qn("w:rFonts"))
                            if rFonts is not None:
                                rPr.remove(rFonts)

    logger.info("Step 2 — styles: normalized %d, zeroed preamble spacing, fixed %d list indents", style_changes, list_fixes)

    # ── Step 3: Insert section breaks at known positions ─────────────────

    def _make_sect_pr(break_type: str, cols: int):
        sect_pr = etree.SubElement(etree.Element("dummy"), qn("w:sectPr"))
        btype = etree.SubElement(sect_pr, qn("w:type"))
        btype.set(qn("w:val"), break_type)
        pg_sz = etree.SubElement(sect_pr, qn("w:pgSz"))
        pg_sz.set(qn("w:w"), "12240")
        pg_sz.set(qn("w:h"), "15840")
        pg_mar = etree.SubElement(sect_pr, qn("w:pgMar"))
        for attr in ["top", "bottom", "left", "right"]:
            pg_mar.set(qn(f"w:{attr}"), _MARGIN_TWIPS)
        pg_mar.set(qn("w:header"), "720")
        pg_mar.set(qn("w:footer"), "720")
        cols_el = etree.SubElement(sect_pr, qn("w:cols"))
        if cols > 1:
            cols_el.set(qn("w:num"), str(cols))
            cols_el.set(qn("w:space"), "480")
        else:
            cols_el.set(qn("w:space"), "720")
        return sect_pr

    def _ensure_pPr(para_el):
        pPr = para_el.find(qn("w:pPr"))
        if pPr is None:
            pPr = etree.SubElement(para_el, qn("w:pPr"))
            para_el.insert(0, pPr)
        return pPr

    # Break 1: End heading (1-col) → newPage before city policies
    _ensure_pPr(paras[idx_city_policies - 1]._element).append(
        _make_sect_pr("nextPage", 1)
    )
    # Break 2: End policies (2-col) → continuous before proposals
    if idx_proposals is not None:
        _ensure_pPr(paras[idx_proposals - 1]._element).append(
            _make_sect_pr("continuous", 2)
        )
    # Break 3: End proposals (2-col) → continuous before projects
    _ensure_pPr(paras[idx_first_project - 1]._element).append(
        _make_sect_pr("continuous", 2)
    )

    # Fix body (final) section: margins + 1-col
    body_sect = doc.sections[-1]._sectPr
    pg_mar = body_sect.find(qn("w:pgMar"))
    if pg_mar is not None:
        for attr in ["top", "bottom", "left", "right"]:
            pg_mar.set(qn(f"w:{attr}"), _MARGIN_TWIPS)
    cols_el = body_sect.find(qn("w:cols"))
    if cols_el is not None and cols_el.get(qn("w:num")):
        del cols_el.attrib[qn("w:num")]

    logger.info("Step 3 — section breaks: 3 inserted + body margins fixed")

    # ── Step 4: Insert transition + spacing paragraphs ───────────────────
    # After this step, paragraph indices from step 1 are INVALID.

    def _make_blank_para():
        return etree.fromstring(
            '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )

    def _make_blank_with_break(break_type, cols):
        el = _make_blank_para()
        pPr = etree.SubElement(el, qn("w:pPr"))
        el.insert(0, pPr)
        pPr.append(_make_sect_pr(break_type, cols))
        return el

    # Two transition blank paragraphs before first project (sections 3-4)
    first_proj_el = paras[idx_first_project]._element
    for _ in range(2):
        first_proj_el.addprevious(_make_blank_with_break("continuous", 2))

    # 3 blank paragraphs between each project block (backwards to preserve indices)
    for idx in reversed(project_header_indices[1:]):
        el = paras[idx]._element
        el.addprevious(deepcopy(_make_blank_para()))

    logger.info("Step 4 — inserted 2 transition + %d×3 inter-project blanks", len(project_header_indices) - 1)

    # ── Step 5: Save and reopen fresh ────────────────────────────────────
    # All XML mutations are done.  Save, then reopen to get a clean
    # paragraph list with proper python-docx wrappers.

    doc.save(str(docx_path))
    doc = Document(str(docx_path))
    paras = doc.paragraphs

    # Re-find first project (indices shifted)
    idx_first_project = None
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if t.startswith("(") and ":" in t and any(k in t for k in ["Artwork", "Total", "Install"]):
            idx_first_project = i
            break

    if idx_first_project is None:
        logger.warning("Lost first project anchor after insertions")
        doc.save(str(docx_path))
        return

    logger.info("Step 5 — reopened: %d paragraphs, first_project at %d", len(paras), idx_first_project)

    # ── Step 6: Apply spacing and page breaks on clean list ──────────────

    # Normalize project paragraph spacing to match original:
    #   Heading 3: sa=None sb=None ls=None (inherit from style)
    #   Body Normal: sa=50800 sb=None ls=1.15
    # Word COM imports explicit sa/sb from HTML that the original doesn't have.
    for i in range(idx_first_project, len(paras)):
        p = paras[i]
        style_name = p.style.name if p.style else ""
        pf = p.paragraph_format
        pf.space_before = None  # clear Word COM's explicit 0 → inherit

        if "Heading" in style_name:
            pf.space_after = None  # clear Word COM's 25400 → inherit
        else:
            if pf.space_after is None or pf.space_after == 0:
                pf.space_after = 50800  # 4pt in EMU
            if pf.line_spacing is None:
                pf.line_spacing = 1.15

    # Page-break optimization for project blocks
    # Disabled: the original .docx has 0 page_break_before and relies on natural
    # flow.  With spacing now matching the original, forced breaks are more
    # disruptive than helpful.  Re-enable if the user wants explicit control.
    # _apply_project_page_breaks(doc, idx_first_project)

    # Convert character shading to native Word highlight (editable in Word)
    # Word COM imports CSS background-color as w:shd (character shading) which
    # can't be toggled with Word's highlight button.  Native w:highlight can.
    _SHD_TO_HIGHLIGHT = {
        "FFFF00": "yellow",
        "00FF00": "green",
        "00FFFF": "cyan",
        "FF00FF": "magenta",
        "FF0000": "red",
        "0000FF": "blue",
    }
    highlight_fixes = 0
    for p in paras:
        for r in p.runs:
            rPr = r._element.find(qn("w:rPr"))
            if rPr is None:
                continue
            shd = rPr.find(qn("w:shd"))
            if shd is None:
                continue
            fill = (shd.get(qn("w:fill")) or "").upper()
            hl_name = _SHD_TO_HIGHLIGHT.get(fill)
            if hl_name:
                rPr.remove(shd)
                hl_el = etree.SubElement(rPr, qn("w:highlight"))
                hl_el.set(qn("w:val"), hl_name)
                highlight_fixes += 1

    if highlight_fixes:
        logger.info("Step 6b — converted %d shading runs to native highlight", highlight_fixes)

    # ── Step 6.5: Embed SDTs for round-trip parsing ─────────────────────

    if uid_map:
        _embed_sdt_tags(doc, paras, uid_map)

    # ── Step 7: Save ─────────────────────────────────────────────────────

    doc.save(str(docx_path))
    total_paras = len(doc.element.body.findall(f".//{qn('w:p')}"))
    logger.info("Step 7 — saved: %d paras, %d sections", total_paras, len(doc.sections))


# ---------------------------------------------------------------------------
# Page-break optimization for project blocks
# ---------------------------------------------------------------------------

_USABLE_PAGE_HEIGHT = 8_618_220  # 11in page - 2×0.79in margins (EMU)
_LINE_HEIGHT_EMU = 160_655       # 11pt × 1.15 line spacing
_SPACE_AFTER_EMU = 50_800        # 4pt space after
_BLANK_PARA_HEIGHT = 152_400     # ~12pt empty paragraph
_CHARS_PER_LINE = 85             # 11pt Calibri at ~6.9in width
_SPILL_THRESHOLD = 2.0 / 5.0    # < 2/5 overflow → push to next page


def _estimate_para_height(para) -> int:
    text = (para.text or "").strip()
    if not text:
        return _BLANK_PARA_HEIGHT
    num_lines = max(1, -(-len(text) // _CHARS_PER_LINE))
    return num_lines * _LINE_HEIGHT_EMU + _SPACE_AFTER_EMU


def _is_project_header(text: str) -> bool:
    return (
        text.startswith("(")
        and ":" in text
        and any(kw in text for kw in ("Artwork", "Total", "Install"))
    )


def _apply_project_page_breaks(doc, first_project_idx: int) -> None:
    paras = doc.paragraphs
    total = len(paras)

    project_starts: list[int] = []
    for i in range(first_project_idx, total):
        if _is_project_header((paras[i].text or "").strip()):
            project_starts.append(i)

    if len(project_starts) < 2:
        return

    breaks_added = 0
    page_pos = 0

    for j, start_idx in enumerate(project_starts):
        end_idx = project_starts[j + 1] if j + 1 < len(project_starts) else total
        block_height = sum(_estimate_para_height(paras[i]) for i in range(start_idx, end_idx))

        remaining = _USABLE_PAGE_HEIGHT - page_pos

        if block_height <= remaining:
            page_pos += block_height
        else:
            overflow = block_height - remaining
            ratio = overflow / block_height if block_height > 0 else 0

            if ratio < _SPILL_THRESHOLD:
                paras[start_idx].paragraph_format.page_break_before = True
                breaks_added += 1
                page_pos = block_height
            else:
                page_pos = overflow

        while page_pos > _USABLE_PAGE_HEIGHT:
            page_pos -= _USABLE_PAGE_HEIGHT

    if breaks_added:
        logger.info("Page breaks: added %d to prevent small overflows", breaks_added)


# ---------------------------------------------------------------------------
# SDT (Structured Document Tag) embedding for round-trip parsing
#
# SDTs are Word content controls — invisible to the user, machine-readable.
# We embed them at generation time so the parser can read them back
# deterministically, without heuristic section/label detection.
# ---------------------------------------------------------------------------

def _create_sdt(tag_val: str, alias: str | None = None):
    """Create a w:sdt XML element with tag and optional alias."""
    from docx.oxml.ns import qn
    from lxml import etree

    sdt = etree.Element(qn("w:sdt"))
    sdt_pr = etree.SubElement(sdt, qn("w:sdtPr"))
    tag_el = etree.SubElement(sdt_pr, qn("w:tag"))
    tag_el.set(qn("w:val"), tag_val)
    if alias:
        alias_el = etree.SubElement(sdt_pr, qn("w:alias"))
        alias_el.set(qn("w:val"), alias)
    etree.SubElement(sdt, qn("w:sdtContent"))
    return sdt


def _wrap_in_sdt(parent, elements: list, tag_val: str, alias: str | None = None):
    """Wrap XML elements in a w:sdt within parent, preserving position."""
    from docx.oxml.ns import qn

    if not elements:
        return
    first_idx = list(parent).index(elements[0])
    sdt = _create_sdt(tag_val, alias)
    content = sdt.find(qn("w:sdtContent"))
    for el in elements:
        content.append(el)  # lxml removes from parent and adds to content
    parent.insert(first_idx, sdt)


def _detect_bold_label(para_el) -> str | None:
    """Detect a bold section label from the first text-bearing run.

    Returns the lowercase label text (without trailing colon) or None.
    Only checks the first run with text — if it's not bold, returns None.
    """
    from docx.oxml.ns import qn

    for r in para_el.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        is_bold = False
        if rPr is not None:
            b = rPr.find(qn("w:b"))
            if b is not None:
                val = b.get(qn("w:val"))
                is_bold = val is None or val not in ("0", "false")

        t_el = r.find(qn("w:t"))
        text = (t_el.text or "").strip() if t_el is not None else ""
        if not text:
            continue

        if not is_bold:
            return None  # First text-bearing run is not bold — no label
        return text.lower().rstrip(":").strip()

    return None


def _group_paragraphs_by_section(
    para_elements: list,
    section_labels: dict[str, str],
    contact_related: set[str],
) -> list[tuple[str, list]]:
    """Group paragraph elements by detected section label.

    Uses bold-label detection (reliable since we generated the labels).
    Folds contact-related sub-sections (architect, ppap, etc.) into "contacts".

    Returns [(section_name, [elements]), ...] in document order.
    """
    groups: list[tuple[str, list]] = []
    current_section: str | None = None
    current_els: list = []

    for el in para_elements:
        raw_label = _detect_bold_label(el)
        new_section = None

        if raw_label:
            mapped = section_labels.get(raw_label)
            if mapped:
                if mapped in contact_related:
                    mapped = "contacts"
                new_section = mapped

        if new_section and new_section != current_section:
            if current_section is not None and current_els:
                groups.append((current_section, current_els))
            current_section = new_section
            current_els = [el]
        else:
            current_els.append(el)

    if current_section is not None and current_els:
        groups.append((current_section, current_els))

    return groups


def _match_headers_to_uids(
    paras, header_indices: list[int], uid_map: list[dict[str, str]],
) -> list[tuple[int, str]]:
    """Match .docx header paragraph indices to UIDs.

    When counts match, uses positional matching (fast, deterministic).
    When counts differ, falls back to text-similarity matching using
    the header hints extracted from the HTML before Word COM stripped them.

    Returns [(header_para_index, uid), ...] for matched pairs.
    """
    if len(header_indices) == len(uid_map):
        return [(idx, entry["uid"]) for idx, entry in zip(header_indices, uid_map)]

    logger.warning(
        "SDT embedding: %d headers vs %d UIDs — using text matching",
        len(header_indices), len(uid_map),
    )

    # Build lookup: header_hint_lower → uid (from HTML extraction)
    remaining = list(uid_map)
    matched: list[tuple[int, str]] = []

    for idx in header_indices:
        docx_header = (paras[idx].text or "").strip().lower()
        best_entry = None
        best_score = 0

        for entry in remaining:
            hint = entry["header"].lower()
            if not hint:
                continue
            # Exact match
            if docx_header == hint:
                best_entry = entry
                best_score = 3
                break
            # Substring match (one contains the other)
            if hint in docx_header or docx_header in hint:
                if best_score < 2:
                    best_entry = entry
                    best_score = 2
            # Client + project name overlap (first 30 chars)
            elif docx_header[:30] == hint[:30] and best_score < 1:
                best_entry = entry
                best_score = 1

        if best_entry:
            matched.append((idx, best_entry["uid"]))
            remaining.remove(best_entry)
        else:
            logger.warning("SDT: no UID match for header at para %d: %s", idx, docx_header[:60])

    if remaining:
        logger.warning("SDT: %d UIDs unmatched: %s", len(remaining),
                       [e["uid"] for e in remaining])

    return matched


def _embed_sdt_tags(doc, paras, uid_map: list[dict[str, str]]) -> None:
    """Wrap project blocks and their sections in Word SDTs.

    Each project block gets a project-level SDT with tag "bfa:project:uid=...".
    Within each project, sections are wrapped with "bfa:section:..." tags.
    SDTs are invisible in Word but survive editing and enable deterministic parsing.
    """
    from docx.oxml.ns import qn
    from .pipeline.config import SECTION_LABELS
    from .pipeline.importer import CONTACT_RELATED_SECTIONS

    body = doc.element.body

    # Find all project header indices
    header_indices = []
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if _is_project_header(t):
            header_indices.append(i)

    if not header_indices:
        logger.warning("SDT embedding: no project headers found")
        return

    # Match headers to UIDs (positional when counts match, text-similarity otherwise)
    matched = _match_headers_to_uids(paras, header_indices, uid_map)
    if not matched:
        logger.warning("SDT embedding: no headers matched to UIDs — skipping")
        return

    # Build end-index lookup: header_idx → next header or end of doc
    header_set = set(header_indices)
    sorted_headers = sorted(header_set)

    # Process in reverse to preserve element indices
    for header_idx, uid in reversed(matched):
        # Find end of this project block
        pos_in_sorted = sorted_headers.index(header_idx)
        end = sorted_headers[pos_in_sorted + 1] if pos_in_sorted + 1 < len(sorted_headers) else len(paras)

        # Trim trailing blanks and ON HOLD separator from block
        while end > header_idx + 1:
            trailing_text = (paras[end - 1].text or "").strip()
            if not trailing_text or "ON HOLD" in trailing_text:
                end -= 1
            else:
                break

        header_text = (paras[header_idx].text or "").strip()
        block_els = [paras[i]._element for i in range(header_idx, end)]
        if not block_els:
            continue

        # --- Project-level SDT ---
        proj_sdt = _create_sdt(
            f"bfa:project:uid={uid}",
            alias=header_text[:80],
        )
        proj_content = proj_sdt.find(qn("w:sdtContent"))

        first_pos = list(body).index(block_els[0])
        for el in block_els:
            proj_content.append(el)
        body.insert(first_pos, proj_sdt)

        # --- Section-level SDTs within project ---
        inner_paras = list(proj_content.findall(qn("w:p")))
        if len(inner_paras) < 2:
            continue  # Just header, no body paragraphs

        body_paras = inner_paras[1:]  # Skip header paragraph
        section_groups = _group_paragraphs_by_section(
            body_paras, SECTION_LABELS, CONTACT_RELATED_SECTIONS,
        )

        for sec_name, sec_elements in reversed(section_groups):
            if sec_elements:
                _wrap_in_sdt(proj_content, sec_elements, f"bfa:section:{sec_name}")

    sdt_count = len(body.findall(f".//{qn('w:sdt')}"))
    logger.info(
        "Step 6.5 — embedded %d SDTs (%d of %d projects matched)",
        sdt_count, len(matched), len(uid_map),
    )


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def render_todo_docx(
    projects: list[dict] | None = None,
    output_path: Path | None = None,
) -> Path:
    """Generate .docx from the pipeline's rendered HTML via Word COM."""
    from .pipeline import config

    site_dir = config.SITE_DIR
    pasteable_path = site_dir / "index_pasteable.html"

    if not pasteable_path.exists():
        raise FileNotFoundError(
            f"Rendered HTML not found at {pasteable_path}. Run the render pipeline first."
        )

    raw_html = pasteable_path.read_text(encoding="utf-8")
    cleaned_html, uid_map = _prepare_html_for_word(raw_html)

    tmp_dir = tempfile.mkdtemp(prefix="bfa_docx_")
    tmp_html = Path(tmp_dir) / "todo_prepared.html"
    tmp_html.write_text(cleaned_html, encoding="utf-8")

    if output_path is None:
        from datetime import date
        today = date.today().strftime("%Y_%m_%d")
        filename = f"{today}_Ballard Fine Art To Do.docx"
        out_dir = config.OUTPUT_DIR
        out_dir.mkdir(parents=True, exist_ok=True)
        output_path = out_dir / filename

    _html_to_docx_via_word(tmp_html, output_path)
    _postprocess_docx(output_path, uid_map=uid_map)

    try:
        tmp_html.unlink()
        Path(tmp_dir).rmdir()
    except Exception:
        pass

    # Also produce PDF via Word COM
    pdf_path = output_path.with_suffix(".docx.pdf")
    _docx_to_pdf(output_path, pdf_path)

    logger.info("Rendered To Do .docx: %s (%d bytes)", output_path, output_path.stat().st_size)
    logger.info("Rendered To Do .pdf: %s (%d bytes)", pdf_path, pdf_path.stat().st_size)
    return output_path


async def upload_to_drive(docx_path: Path, file_id: str) -> dict:
    """Upload .docx to Google Drive, replacing existing file content."""
    from autohelper.modules.bfa_todo.google_client import GoogleDocsClient

    client = GoogleDocsClient()
    token = client._get_token()

    import httpx
    async with httpx.AsyncClient() as http:
        with open(docx_path, "rb") as f:
            r = await http.patch(
                f"https://www.googleapis.com/upload/drive/v3/files/{file_id}",
                headers={"Authorization": f"Bearer {token}"},
                content=f.read(),
                params={"uploadType": "media"},
                timeout=120.0,
            )

    if r.status_code >= 400:
        return {"error": f"Drive upload failed: {r.status_code} {r.text[:200]}"}

    result = r.json()
    logger.info("Uploaded %s to Drive file %s", docx_path.name, file_id)
    return {"file_id": result.get("id"), "name": result.get("name")}
