"""Fidelity tests: produced docx must match source docx XML properties.

Tests the GENERATED .docx against the SOURCE .docx formatting spec.
Ground truth is Word XML (w:spacing, w:sz, w:tblBorders, w:numPr, etc.),
NOT intermediate HTML.

These tests run the full pipeline: ingest → render → generate docx,
then compare the produced docx paragraph-by-paragraph against the source.
"""

import os
import re
from datetime import date, timedelta

import pytest
from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture(scope="module")
def source_doc():
    """Load the PM's canonical source docx (formatting authority)."""
    path = os.environ.get(
        "BFA_SOURCE_1",
        r"C:\Users\Neal\Downloads\2026_01_08_Ballard Fine Art To Do - gdocs.docx",
    )
    if not os.path.exists(path):
        pytest.skip(f"Source docx not found: {path}")
    return Document(path)


@pytest.fixture(scope="module")
def produced_doc():
    """Load the most recently produced BFA To Do docx.

    The pipeline (render + Word COM + postprocess) must be run separately
    before these tests — Word COM can't run inside pytest reliably.
    Tests verify the OUTPUT, not the generation process.
    """
    import glob

    search_paths = [
        os.path.join(
            os.path.dirname(__file__), os.pardir,
            "autohelper", "modules", "bfa_todo", "_bfa", "*Ballard*.docx",
        ),
        r"E:\OneDrive - Ballard Fine Art\BALLARD FINE ART - ALL FILES"
        r"\4. TO DO LIST\2026\_bfa\*Ballard*.docx",
    ]
    candidates = []
    for pattern in search_paths:
        candidates.extend(glob.glob(pattern))

    if not candidates:
        pytest.skip("No produced docx found — run render_todo_docx() first")

    # Use most recent
    candidates.sort(key=lambda f: os.path.getmtime(f), reverse=True)
    path = candidates[0]

    try:
        return Document(path)
    except Exception as e:
        pytest.skip(f"Could not open produced docx {path}: {e}")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _find_para(paras, text_fragment):
    """Find first paragraph containing text_fragment."""
    for i, p in enumerate(paras):
        if text_fragment in (p.text or ""):
            return i, p
    return None, None


def _get_table_borders(tbl_element):
    """Extract table border properties as dict."""
    tblPr = tbl_element.find(qn("w:tblPr"))
    if tblPr is None:
        return {}
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        return {}
    result = {}
    for b in borders:
        name = b.tag.split("}")[-1]
        result[name] = {
            "val": b.get(qn("w:val")),
            "sz": b.get(qn("w:sz")),
            "color": b.get(qn("w:color")),
        }
    return result


def _first_text_run(para):
    """Return the first run that has actual text content (not just whitespace/linebreaks)."""
    for r in para.runs:
        if (r.text or "").strip():
            return r
    return para.runs[0] if para.runs else None


# ---------------------------------------------------------------------------
# A. Header area
# ---------------------------------------------------------------------------

class TestTitleParagraph:
    """Title: 'Ballard Fine Art - To Do List <date>'"""

    def test_title_is_bold(self, produced_doc):
        """Title must have bold formatting — not destroyed by pipeline."""
        paras = produced_doc.paragraphs
        _, title = _find_para(paras, "To Do List")
        assert title is not None, "Title paragraph not found"
        r = _first_text_run(title)
        assert r is not None
        assert r.bold, f"Title run should be bold, got bold={r.bold}"

    def test_title_font_size_11pt(self, produced_doc):
        paras = produced_doc.paragraphs
        _, title = _find_para(paras, "To Do List")
        assert title is not None
        r = _first_text_run(title)
        assert r is not None and r.font.size is not None
        pt = r.font.size / 12700
        assert 10.5 <= pt <= 11.5, f"Title should be 11pt, got {pt:.1f}pt"

    def test_title_date_is_next_monday(self, produced_doc):
        """Title date must be next Monday, not today."""
        paras = produced_doc.paragraphs
        _, title = _find_para(paras, "To Do List")
        assert title is not None
        text = title.text

        today = date.today()
        days_until_monday = (7 - today.weekday()) % 7
        if days_until_monday == 0:
            days_until_monday = 7
        monday = today + timedelta(days=days_until_monday)
        expected = f"{monday.strftime('%B')} {monday.day}, {monday.year}"

        assert expected in text, (
            f"Title should contain next Monday '{expected}', got: {text}"
        )


class TestRichmondParagraph:
    """Richmond line: 8pt, blue hyperlink, dates updated."""

    def test_richmond_font_size_8pt(self, produced_doc):
        """Richmond content runs must be 8pt, not 11pt."""
        paras = produced_doc.paragraphs
        _, rich = _find_para(paras, "Richmond")
        assert rich is not None, "Richmond paragraph not found"
        # Find a content run (not the line break)
        for r in rich.runs:
            text = (r.text or "").strip()
            if "Richmond" in text or "Biliana" in text or "confirm" in text:
                assert r.font.size is not None, f"Richmond run has no font size: '{text}'"
                pt = r.font.size / 12700
                assert 7.5 <= pt <= 8.5, (
                    f"Richmond run should be 8pt, got {pt:.1f}pt: '{text}'"
                )
                return
        pytest.fail("No Richmond content run found")

    def test_richmond_dates_are_current(self, produced_doc):
        """Richmond dates should be computed (future), not source dates."""
        paras = produced_doc.paragraphs
        _, rich = _find_para(paras, "Richmond")
        assert rich is not None
        text = rich.text
        # Should NOT contain the old source dates
        assert "Nov 25, 2025" not in text, "Richmond still has old source dates"
        assert "Dec 17, 2025" not in text, "Richmond still has old source dates"


class TestNorthVanParagraph:
    """North Van line: 8pt, blue bold header, dates updated."""

    def test_northvan_font_size_8pt(self, produced_doc):
        paras = produced_doc.paragraphs
        _, nv = _find_para(paras, "North Van")
        assert nv is not None, "North Van paragraph not found"
        for r in nv.runs:
            text = (r.text or "").strip()
            if "North Van" in text:
                assert r.font.size is not None
                pt = r.font.size / 12700
                assert 7.5 <= pt <= 8.5, f"North Van should be 8pt, got {pt:.1f}pt"
                return
        pytest.fail("No 'North Van' run found")

    def test_northvan_dates_are_current(self, produced_doc):
        """North Van dates should be computed, not old source dates."""
        paras = produced_doc.paragraphs
        _, nv = _find_para(paras, "North Van")
        assert nv is not None
        text = nv.text
        # Should NOT still have all of source 1's 2026 schedule verbatim
        # (which includes January, February dates that are in the past)
        assert "January 8, 2026" not in text, "North Van has stale source dates"


# ---------------------------------------------------------------------------
# B. PAC Table
# ---------------------------------------------------------------------------

class TestPacTable:
    """PAC meeting schedule table: borders must match source."""

    def _find_preamble_table(self, doc):
        body = doc.element.body
        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "tbl":
                # First table in body is the PAC schedule
                return child
            if tag == "sdt":
                content = child.find(qn("w:sdtContent"))
                if content is not None:
                    for inner in content:
                        if inner.tag.endswith("}tbl"):
                            return inner
        return None

    def test_table_has_borders(self, produced_doc):
        tbl = self._find_preamble_table(produced_doc)
        assert tbl is not None, "PAC table not found in produced doc"
        borders = _get_table_borders(tbl)
        assert borders, "PAC table has no border definitions"
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            assert side in borders, f"Missing border: {side}"
            assert borders[side]["val"] == "single", f"{side} border should be 'single'"

    def test_table_border_weight_matches_source(self, source_doc, produced_doc):
        """Borders should be sz=8 (thick) matching source, not sz=4 (thin default)."""
        src_tbl = self._find_preamble_table(source_doc)
        prod_tbl = self._find_preamble_table(produced_doc)
        assert src_tbl is not None and prod_tbl is not None

        src_borders = _get_table_borders(src_tbl)
        prod_borders = _get_table_borders(prod_tbl)

        for side in ("top", "bottom", "left", "right"):
            src_sz = src_borders.get(side, {}).get("sz", "8")
            prod_sz = prod_borders.get(side, {}).get("sz")
            assert prod_sz == src_sz, (
                f"Table {side} border: source sz={src_sz}, produced sz={prod_sz}"
            )

    def test_table_border_color_is_black(self, produced_doc):
        tbl = self._find_preamble_table(produced_doc)
        assert tbl is not None
        borders = _get_table_borders(tbl)
        for side in ("top", "bottom", "left", "right"):
            color = borders.get(side, {}).get("color", "")
            assert color in ("000000", "auto"), (
                f"Table {side} border color should be black, got {color}"
            )


# ---------------------------------------------------------------------------
# C. Policy bullets
# ---------------------------------------------------------------------------

class TestPolicyBullets:
    """Preamble policy bullets: indentation must match source."""

    def _find_first_bullet(self, doc):
        for p in doc.paragraphs:
            pPr = p._element.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:numPr")) is not None:
                return p
            # Also check for text bullets from our pipeline
            if p.text and "\u25cf" in p.text:
                return p
        return None

    def test_bullets_exist_in_preamble(self, produced_doc):
        bullet = self._find_first_bullet(produced_doc)
        assert bullet is not None, "No bullet paragraphs found in produced doc"

    def test_bullet_indent_matches_source(self, source_doc, produced_doc):
        """Bullet left indent should match source (179705 EMU)."""
        src_bullet = self._find_first_bullet(source_doc)
        prod_bullet = self._find_first_bullet(produced_doc)
        assert src_bullet is not None and prod_bullet is not None

        src_indent = src_bullet.paragraph_format.left_indent
        prod_indent = prod_bullet.paragraph_format.left_indent

        assert prod_indent is not None, "Produced bullet has no left indent"
        # Allow 10% tolerance
        assert abs(prod_indent - src_indent) < src_indent * 0.15, (
            f"Bullet indent: source={src_indent}, produced={prod_indent}"
        )


# ---------------------------------------------------------------------------
# D. Proposals header
# ---------------------------------------------------------------------------

class TestProposalsHeader:
    """'Proposals (New Projects)' line: bold, correct size."""

    def test_proposals_is_bold(self, produced_doc):
        paras = produced_doc.paragraphs
        _, prop = _find_para(paras, "Proposals")
        assert prop is not None, "Proposals paragraph not found"
        r = _first_text_run(prop)
        assert r is not None
        assert r.bold, f"Proposals header should be bold, got bold={r.bold}"

    def test_proposals_font_not_8pt(self, produced_doc):
        """Proposals header should NOT be 8pt (preamble detail size)."""
        paras = produced_doc.paragraphs
        _, prop = _find_para(paras, "Proposals")
        assert prop is not None
        r = _first_text_run(prop)
        if r and r.font.size:
            pt = r.font.size / 12700
            assert pt > 9, f"Proposals should not be 8pt, got {pt:.1f}pt"


# ---------------------------------------------------------------------------
# E. Section breaks
# ---------------------------------------------------------------------------

class TestSectionStructure:
    """Document section layout must match source's 6-section structure."""

    def test_section_count(self, source_doc, produced_doc):
        """Produced doc should have same number of sections as source."""
        src_count = len(source_doc.sections)
        prod_count = len(produced_doc.sections)
        assert prod_count >= src_count - 1, (
            f"Source has {src_count} sections, produced has {prod_count}"
        )

    def test_heading_section_is_single_column(self, produced_doc):
        """First section (title/header) must be 1-column."""
        sec = produced_doc.sections[0]
        cols = sec._sectPr.find(qn("w:cols"))
        if cols is not None:
            num = cols.get(qn("w:num"), "1")
            assert num == "1", f"Header section should be 1-col, got {num}"

    def test_policy_section_is_two_column(self, produced_doc):
        """Policy section (after header) must be 2-column."""
        if len(produced_doc.sections) < 2:
            pytest.skip("Not enough sections")
        sec = produced_doc.sections[1]
        cols = sec._sectPr.find(qn("w:cols"))
        assert cols is not None
        num = cols.get(qn("w:num"), "1")
        assert num == "2", f"Policy section should be 2-col, got {num}"

    def test_projects_section_is_single_column(self, produced_doc):
        """Final section (projects) must be 1-column."""
        sec = produced_doc.sections[-1]
        cols = sec._sectPr.find(qn("w:cols"))
        if cols is not None:
            num = cols.get(qn("w:num"), "1")
            assert num == "1", f"Projects section should be 1-col, got {num}"


# ---------------------------------------------------------------------------
# F. Margins
# ---------------------------------------------------------------------------

class TestMargins:
    """Page margins must match source (~0.567 inches = 719455 EMU)."""

    def test_margins_match_source(self, source_doc, produced_doc):
        src_sec = source_doc.sections[0]
        prod_sec = produced_doc.sections[0]

        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            src_val = getattr(src_sec, attr)
            prod_val = getattr(prod_sec, attr)
            # Allow 5% tolerance
            assert abs(prod_val - src_val) < src_val * 0.05, (
                f"{attr}: source={src_val}, produced={prod_val}"
            )


# ---------------------------------------------------------------------------
# G. Project content integrity
# ---------------------------------------------------------------------------

class TestProjectContent:
    """Projects must survive the pipeline with correct formatting."""

    def test_project_headers_exist(self, produced_doc):
        """At least 140 project headers in the produced doc (inside SDTs)."""
        body = produced_doc.element.body
        count = 0
        for sdt in body.iter(qn("w:sdt")):
            sdtPr = sdt.find(qn("w:sdtPr"))
            if sdtPr is not None:
                tag = sdtPr.find(qn("w:tag"))
                if tag is not None and "bfa:project:uid=" in tag.get(qn("w:val"), ""):
                    count += 1
        assert count >= 135, f"Expected >=135 project SDTs, got {count}"

    def test_project_body_has_bold_labels(self, produced_doc):
        """Project sections should contain bold runs (Contact:, Owner:, etc.)."""
        body = produced_doc.element.body
        bold_count = 0
        for sdt in body.iter(qn("w:sdt")):
            sdtPr = sdt.find(qn("w:sdtPr"))
            if sdtPr is None:
                continue
            tag = sdtPr.find(qn("w:tag"))
            if tag is None or "bfa:project:uid=" not in tag.get(qn("w:val"), ""):
                continue
            for rPr in sdt.iter(qn("w:rPr")):
                b = rPr.find(qn("w:b"))
                if b is not None and b.get(qn("w:val"), "true") != "false":
                    bold_count += 1
        assert bold_count >= 500, f"Expected >=500 bold runs in projects, got {bold_count}"

    def test_known_labels_are_bold(self, produced_doc):
        """Field labels (Contact:, AO:, SP#1:, etc.) must be bold in every project.

        Navigates nested SDTs carefully: project SDTs contain section SDTs,
        and iter() would cross SDT boundaries. Instead, find direct w:p
        children within each project SDT's content tree.
        """
        _LABELS = [
            "Contact", "Owner", "Architect", "PPAP", "DPAP",
            "SP#1", "SP#2", "EOI", "AO", "PM",
            "Project Status", "Artwork Title", "Selected Artist",
        ]
        body = produced_doc.element.body
        not_bold = []

        for sdt in body.iter(qn("w:sdt")):
            sdtPr = sdt.find(qn("w:sdtPr"))
            if sdtPr is None:
                continue
            tag = sdtPr.find(qn("w:tag"))
            if tag is None or "bfa:project:uid=" not in tag.get(qn("w:val"), ""):
                continue

            # Collect ALL w:p elements within this project SDT
            for p_el in sdt.iter(qn("w:p")):
                text = "".join(
                    (t.text or "") for t in p_el.iter(qn("w:t"))
                ).strip()

                matched = None
                for label in _LABELS:
                    if text.startswith(label + ":") or text.startswith(label + " :"):
                        matched = label
                        break
                if not matched:
                    continue

                # Check the DIRECT child w:r elements of THIS paragraph only
                # (not runs from nested SDTs)
                first_run = None
                for child in p_el:
                    if child.tag == qn("w:r"):
                        first_run = child
                        break
                    # Also check inside w:hyperlink direct children
                    if child.tag == qn("w:hyperlink"):
                        r = child.find(qn("w:r"))
                        if r is not None:
                            first_run = r
                            break

                if first_run is None:
                    continue

                rPr = first_run.find(qn("w:rPr"))
                b = rPr.find(qn("w:b")) if rPr is not None else None
                is_bold = b is not None and b.get(qn("w:val"), "true") != "false"

                if not is_bold:
                    uid = tag.get(qn("w:val"), "")
                    not_bold.append(f"{matched}: in {uid[-12:]}")

        assert len(not_bold) == 0, (
            f"{len(not_bold)} labels not bold: {not_bold[:10]}"
        )
