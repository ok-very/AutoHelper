"""
Create Jinja2-marked-up .docx templates from precedent files.

Copies each precedent, replaces hardcoded values with {{merge_fields}},
and adds a draft generation footer to each template.

Run: mise exec -- python scripts/create_docx_templates.py
"""

import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PRECEDENTS = Path(
    r"C:\Users\silen\OneDrive - Ballard Fine Art"
    r"\BALLARD FINE ART - ALL FILES\6. PA PROCEDURES\_bfa\precedents"
)
TEMPLATES = PRECEDENTS.parent / "templates"
TEMPLATES.mkdir(exist_ok=True)


def add_draft_footer(doc: Document, template_name: str) -> None:
    """Add a draft generation footer with Jinja2 timestamp + version fields."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        # Clear existing footer paragraphs beyond the first
        # (keep structure, we'll use the first paragraph)
        if footer.paragraphs:
            p = footer.paragraphs[0]
        else:
            p = footer.add_paragraph()

        # If footer already has content, add a new paragraph instead
        if p.text.strip():
            p = footer.add_paragraph()

        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("{{ generated_at }}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.italic = True


def replace_in_paragraphs(doc: Document, replacements: dict[str, str]) -> int:
    """Replace text in paragraphs, preserving formatting of the first run."""
    count = 0
    all_paragraphs = list(doc.paragraphs)
    # Also include paragraphs inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    for p in all_paragraphs:
        full_text = p.text
        for old, new in replacements.items():
            if old in full_text:
                # Rebuild runs: clear all, put merged text in first run
                # Preserve first run's formatting
                if p.runs:
                    fmt = p.runs[0].font
                    font_name = fmt.name
                    font_size = fmt.size
                    font_bold = fmt.bold
                    font_italic = fmt.italic
                    font_color = fmt.color.rgb if fmt.color and fmt.color.rgb else None
                else:
                    font_name = font_size = font_bold = font_italic = font_color = None

                new_text = full_text.replace(old, new)
                full_text = new_text

                # Clear existing runs
                for run in p.runs:
                    run.text = ""
                if p.runs:
                    p.runs[0].text = new_text
                    # Restore formatting
                    if font_name:
                        p.runs[0].font.name = font_name
                    if font_size:
                        p.runs[0].font.size = font_size
                    if font_bold is not None:
                        p.runs[0].font.bold = font_bold
                    if font_italic is not None:
                        p.runs[0].font.italic = font_italic
                    if font_color:
                        p.runs[0].font.color.rgb = font_color
                else:
                    run = p.add_run(new_text)

                count += 1
    return count


# ── Artwork Acceptance ──────────────────────────────────────────

def template_artwork_acceptance():
    src = PRECEDENTS / "2025_Artwork Acceptance Template.docx"
    dst = TEMPLATES / "artwork_acceptance.docx"
    shutil.copy2(src, dst)
    doc = Document(str(dst))

    replacements = {
        # Project / artwork
        "The Measure": "{{ artwork_title }}",
        "5977 Wilson Ave, Burnaby B.C V5H 4R8": "{{ site_address }}",
        "5977 Wilson Ave, Burnaby B.C. V5H 4R8": "{{ site_address }}",
        "February 12th, 2025": "{{ install_date }}",
        # Artist
        "A. Kwade GmbH": "{{ artist_name }}",
        "Reinbeckstraße 22 12459 Berlin,": "{{ artist_address_line1 }}",
        "Reinbeckstra\u00dfe 22 12459 Berlin,": "{{ artist_address_line1 }}",
        "Germany Phone: +49 30 6408 1075": "{{ artist_address_line2 }}",
        "Germany Phone: +49 30 6408 1075 ": "{{ artist_address_line2 }}",
        "sonia@alicjakwade.com": "{{ artist_email }}",
        # Owner
        "Bosa Properties (Wilson) Limited Partnership": "{{ developer_name }}",
        "1100 – 838 West Hastings St,": "{{ developer_address_line1 }}",
        "1100 \u2013 838 West Hastings St,": "{{ developer_address_line1 }}",
        "Vancouver, B.C. V66 0A6": "{{ developer_address_line2 }}",
        "(604) 229 1363": "{{ developer_phone }}",
        "Philippe Lew, Development Mana": "{{ developer_contact_name }}, {{ developer_contact_title }}",
    }

    n = replace_in_paragraphs(doc, replacements)
    add_draft_footer(doc, "Artwork Acceptance")
    doc.save(str(dst))
    print(f"  artwork_acceptance.docx — {n} replacements")


# ── Transfer of Title ───────────────────────────────────────────

def template_transfer_of_title():
    src = PRECEDENTS / "Transfer of Title Template 2025.docx"
    dst = TEMPLATES / "transfer_of_title.docx"
    shutil.copy2(src, dst)
    doc = Document(str(dst))

    replacements = {
        # Project / artwork
        "The Measure": "{{ artwork_title }}",
        "Central Park House": "{{ development_name }}",
        "5977 Wilson Ave": "{{ site_address }}",
        # Developer
        "Bosa Properties (Wilson) LP": "{{ developer_name }}",
        "Bosa Properties (Wilson)LP": "{{ developer_name }}",
        "Bosa Properties (Wilson) Limited Partnership": "{{ developer_name }}",
        "1100 – 838 West Hastings St,": "{{ developer_address_line1 }}",
        "1100 \u2013 838 West Hastings St,": "{{ developer_address_line1 }}",
        "Vancouver, B.C. V66 0A6": "{{ developer_address_line2 }}",
        "(604) 229 1363": "{{ developer_phone }}",
        # Artist
        "A. Kwade GmbH": "{{ artist_name }}",
        "A.Kwade DMBH": "{{ artist_name }}",
        "Reinbeckstraße 22 ": "{{ artist_address_line1 }}",
        "Reinbeckstra\u00dfe 22 ": "{{ artist_address_line1 }}",
        "12459 Berlin": "{{ artist_address_line2 }}",
        "Germany": "{{ artist_country }}",
        "Germany Phone: +49 30 6408 1075": "{{ artist_phone }}",
        # Dates
        "the 14th day of April 2023": "{{ agreement_date }}",
        "Section B.6": "Section {{ agreement_section }}",
    }

    n = replace_in_paragraphs(doc, replacements)
    add_draft_footer(doc, "Transfer of Title")
    doc.save(str(dst))
    print(f"  transfer_of_title.docx — {n} replacements")


# ── Statutory Declaration ───────────────────────────────────────

def template_statutory_declaration():
    src = PRECEDENTS / "2022_04_04_CoV_StatuatoryDeclarationTemplate_01791042_v1.docx"
    dst = TEMPLATES / "statutory_declaration.docx"
    shutil.copy2(src, dst)
    doc = Document(str(dst))

    # This template already uses blanks (___) — replace with merge fields
    replacements = {
        # Declarant info (paragraph 003)
        "I, ___________________________________, of _______________, Vancouver, British Columbia":
            "I, {{ declarant_name }}, of {{ declarant_address }}, {{ city_name }}, British Columbia",
        # Declarant role (paragraph 005)
        "I am the ________________________ of _______________":
            "I am the {{ declarant_title }} of {{ developer_name }}",
        # Artwork + report author (paragraph 007)
        "[insert name of art]": "{{ artwork_title }}",
        "[insert Final Report author\u2019s name]": "{{ final_report_author }}",
        "[insert Final Report author's name]": "{{ final_report_author }}",
        # Registration type
        "[registered / beneficial]": "{{ registration_type }}",
        # Signature block dates
        "DECLARED BEFORE ME at ___________,": "DECLARED BEFORE ME at {{ declaration_city }},",
        "British Columbia, on __________, 202_": "British Columbia, on {{ declaration_date }}",
    }

    n = replace_in_paragraphs(doc, replacements)
    add_draft_footer(doc, "Statutory Declaration")
    doc.save(str(dst))
    print(f"  statutory_declaration.docx — {n} replacements")


# ── Main ────────────────────────────────────────────────────────

if __name__ == "__main__":
    print(f"Precedents: {PRECEDENTS}")
    print(f"Templates:  {TEMPLATES}")
    print()
    template_artwork_acceptance()
    template_transfer_of_title()
    template_statutory_declaration()
    print()
    print("Done. Templates written to:")
    for f in sorted(TEMPLATES.glob("*.docx")):
        print(f"  {f.name}")
