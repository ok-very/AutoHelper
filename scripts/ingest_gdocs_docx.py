"""Ingest BFA To Do List into the DB with rich formatting.

Usage:
  python scripts/ingest_gdocs_docx.py --source path/to/pm_docx.docx

The PM's gdocs docx is the single authority for content and formatting.
Extracts run-level formatting (bold, italic, strikethrough, color, highlight,
hyperlinks) as HTML, matches against DB, and stores with ownership="curated".
Dynamic content (meeting dates) is updated programmatically by the render pipeline.
"""

import argparse
import html as html_mod
import re
import sys
import uuid

from docx import Document
from docx.oxml.ns import qn

parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
parser.add_argument("--source", required=True, help="PM's gdocs docx — formatting and content authority for all sections")
args = parser.parse_args()

from autohelper.config import get_settings
from autohelper.db import init_db
from autohelper.db.migrate import run_migrations

settings = get_settings()
db = init_db(settings.db_path)
run_migrations(db)

from autohelper.modules.bfa_todo.pipeline.repo import BfaProjectRepo
from autohelper.modules.bfa_todo.pipeline.importer import _parse_header_fields
from autohelper.modules.bfa_todo.docx_run_extractor import (
    extract_project_paragraph_html,
    extract_preamble_paragraph_html,
    iter_preamble_body_children,
)

repo = BfaProjectRepo()

# Section label detection — longest match first (sorted at module level)
SECTION_LABELS = {
    "bfa project status": "bfa_phase",
    "bfa next step": "next_steps",
    "project status": "bfa_phase",
    "community advisory": "contacts",
    "selection panel": "contacts",
    "artwork title": "artwork_title",
    "selected artist": "artists",
    "shortlisted": "artists",
    "next steps": "next_steps",
    "next step": "next_steps",
    "bfa phase": "bfa_phase",
    "review of art": "fabrication",
    "final report": "fabrication",
    "installation": "fabrication",
    "milestone": "milestones",
    "governance": "governance",
    "fabricat": "fabrication",
    "landscape": "contacts",
    "architect": "contacts",
    "contacts": "contacts",
    "contact": "contacts",
    "storage": "fabrication",
    "artists": "artists",
    "artist": "artists",
    "nvpaac": "contacts",
    "owner": "contacts",
    "phase": "bfa_phase",
    "ppap": "contacts",
    "dpap": "contacts",
    "team": "contacts",
    "sp#1": "contacts",
    "sp#2": "contacts",
    "eoi": "contacts",
    "ao": "contacts",
}
# Sort longest-first for greedy matching
_SORTED_LABELS = sorted(SECTION_LABELS.items(), key=lambda x: -len(x[0]))


def _normalize_header(text):
    """Normalize header text for matching: dashes, whitespace, case."""
    t = text.replace("\u2013", "-").replace("\u2014", "-")  # en/em dash → hyphen
    t = t.replace("\u00a0", " ")  # nbsp → space
    t = re.sub(r"[\s\t]+", " ", t).strip()
    return t


def detect_section(text):
    t = text.lower().strip()
    for label, sec in _SORTED_LABELS:
        if t.startswith(label):
            return sec
    return None


def is_header_para_text(text):
    """Check if text looks like a project header."""
    if not text:
        return False
    if text.startswith("(") and len(text) > 30:
        if any(kw in text for kw in ["Artwork", "Total", "Install", "Art:", "ON HOLD"]):
            return True
    return False


def is_header_para(p):
    text = p.text.strip()
    if not text:
        return False
    if p.style and p.style.name == "Heading 3":
        return True
    return is_header_para_text(text)


def is_proposals_header(p):
    """Detect 'Proposals (New Projects)' — boundary between preamble and preamble-lists."""
    text = p.text.strip()
    return text.startswith("Proposals") and "New Project" in text


# ── Parse docx ──────────────────────────────────────────────────────────

doc1 = Document(args.source)
paras1 = doc1.paragraphs

# Find paragraph-level boundaries (for project extraction which uses doc.paragraphs)
proposals_para_idx = None
first_project_para_idx = None

for i, p in enumerate(paras1):
    if proposals_para_idx is None and is_proposals_header(p):
        proposals_para_idx = i
    if first_project_para_idx is None and is_header_para(p):
        first_project_para_idx = i
        break

if first_project_para_idx is None:
    print("ERROR: no project headers found")
    raise SystemExit(1)

if proposals_para_idx is None:
    proposals_para_idx = first_project_para_idx

# Find body-child-level boundaries (for preamble extraction which includes SDTs/tables)
body_children = list(doc1.element.body)
proposals_body_idx = None
first_project_body_idx = None

for i, child in enumerate(body_children):
    if child.tag != qn("w:p"):
        continue
    text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
    if proposals_body_idx is None and text.startswith("Proposals") and "New Project" in text:
        proposals_body_idx = i
    if first_project_body_idx is None and is_header_para_text(text):
        first_project_body_idx = i
        break

if proposals_body_idx is None:
    proposals_body_idx = first_project_body_idx or len(body_children)
if first_project_body_idx is None:
    first_project_body_idx = len(body_children)

print(f"Boundaries: proposals_body={proposals_body_idx}, first_project_body={first_project_body_idx}")
print(f"  para-level: proposals={proposals_para_idx}, first_project={first_project_para_idx}, total={len(paras1)}")

# ── 1. Extract preamble (walks body children: paragraphs + SDTs + tables) ──

preamble_texts = []
preamble_htmls = []
for t, h in iter_preamble_body_children(doc1, 0, proposals_body_idx):
    if not t.strip() and "<table" not in h:
        preamble_texts.append("")
        preamble_htmls.append('<p style="margin:0;font-size:8pt;line-height:1.15">&nbsp;</p>')
    else:
        # Tables and paragraphs both become single text lines (replace \n with tab)
        preamble_texts.append(t.replace("\n", "\t"))
        preamble_htmls.append(h)

preamble_entry = {
    "text": "\n".join(preamble_texts),
    "html": "\n".join(preamble_htmls),
}

# ── 2. Extract preamble-lists ────────────────────────────────────────────

preamble_lists_texts = []
preamble_lists_htmls = []
for t, h in iter_preamble_body_children(doc1, proposals_body_idx, first_project_body_idx):
    if not t.strip():
        preamble_lists_texts.append("")
        preamble_lists_htmls.append('<p style="margin:0;font-size:8pt;line-height:1.15">&nbsp;</p>')
    else:
        preamble_lists_texts.append(t.replace("\n", " "))
        preamble_lists_htmls.append(h)

preamble_lists_entry = {
    "text": "\n".join(preamble_lists_texts),
    "html": "\n".join(preamble_lists_htmls),
}

# ── 3. Extract projects ──────────────────────────────────────────────────

docx_projects = []
current = None

for p in paras1[first_project_para_idx:]:
    text = p.text.strip()

    if is_header_para(p):
        if current:
            docx_projects.append(current)
        plain_text, rich_html = extract_project_paragraph_html(p)
        norm_text = _normalize_header(plain_text)
        current = {
            "header_text": norm_text,
            "header_html": rich_html,
            "sections": {},
            "_cur": "contacts",
            "_section_texts": {},
            "_section_htmls": {},
        }
    elif current is not None:
        if not text:
            # Blank paragraph — preserve as spacer in both text and HTML
            sec = current["_cur"]
            current["_section_texts"].setdefault(sec, []).append("")
            current["_section_htmls"].setdefault(sec, []).append("<p>&nbsp;</p>")
            continue

        sec = detect_section(text)
        if sec:
            current["_cur"] = sec

        sec_key = current["_cur"]
        plain_text, rich_html = extract_project_paragraph_html(p)
        current["_section_texts"].setdefault(sec_key, []).append(plain_text)
        current["_section_htmls"].setdefault(sec_key, []).append(rich_html)

if current:
    docx_projects.append(current)

# Finalize sections: join text lines and HTML fragments
for proj in docx_projects:
    for sec_key in set(list(proj["_section_texts"].keys()) + list(proj["_section_htmls"].keys())):
        texts = proj["_section_texts"].get(sec_key, [])
        htmls = proj["_section_htmls"].get(sec_key, [])
        proj["sections"][sec_key] = {
            "text": "\n".join(texts),
            "html": "\n".join(htmls),
        }
    del proj["_cur"]
    del proj["_section_texts"]
    del proj["_section_htmls"]

print(f"Parsed: {len(docx_projects)} projects, preamble, preamble-lists")

# ── Load DB ─────────────────────────────────────────────────────────────

db_entries = repo.get_all(include_archived=False)
db_projects = [e for e in db_entries if e.get("type") == "project"]

# Build lookup index with normalized headers
db_by_header = {}
for e in db_projects:
    h = _normalize_header(e.get("header", {}).get("text", "")).lower()
    db_by_header[h] = e


def find_match(header):
    h = _normalize_header(header).lower()
    if h in db_by_header:
        return db_by_header[h]
    for db_h, e in db_by_header.items():
        if len(h) > 20 and (h[:40] in db_h or db_h[:40] in h):
            return e
    for e in db_projects:
        f = e.get("fields", {})
        cl = f.get("client", "").lower()
        pn = f.get("project_name", "").lower()
        if cl and pn and len(cl) > 2 and len(pn) > 2 and cl in h and pn in h:
            return e
    return None


# ── Store preambles ─────────────────────────────────────────────────────

def _store_preamble(ptype, section_data):
    """Store or update a preamble entry."""
    existing = repo.get_by_type(ptype)
    if existing:
        entry = existing[0].copy()
        entry["sections"]["content"] = section_data
        entry["fields"]["ownership"] = "curated"
        entry["source"] = "gdocs"
        repo.upsert(entry, allow_curated=True)
        return "updated"
    else:
        uid = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"bfa-{ptype}"))
        entry = {
            "uid": uid,
            "slug": f"bfa-{ptype}",
            "fingerprint": f"bfa|{ptype}|",
            "type": ptype,
            "status": "active",
            "source": "gdocs",
            "fields": {"ownership": "curated"},
            "header": {"text": "", "html": ""},
            "sections": {"content": section_data},
            "bfa_phase_display": "",
            "bfa_phase_canonical": "",
            "next_steps": [],
        }
        repo.upsert(entry, allow_curated=True)
        return "created"

r1 = _store_preamble("preamble", preamble_entry)
r2 = _store_preamble("preamble-lists", preamble_lists_entry)
print(f"Preamble: {r1}, Preamble-lists: {r2}")

# ── Match + apply projects ──────────────────────────────────────────────

updated = 0
created = 0
skipped = 0

for proj in docx_projects:
    header_text = proj["header_text"]
    header_html = proj["header_html"]

    if header_text.strip().lower().startswith("project cancel"):
        skipped += 1
        continue

    db_match = find_match(header_text)

    if db_match:
        entry = db_match.copy()
        entry["header"] = {"text": header_text, "html": header_html}
        entry["fields"]["ownership"] = "curated"
        entry["source"] = "gdocs"

        # Update fields from header parse
        try:
            parsed_fields = _parse_header_fields(header_text)
            if parsed_fields:
                for k, v in parsed_fields.items():
                    if v and v != "TBC":
                        entry["fields"][k] = v
        except Exception:
            pass

        # Update sections with rich HTML
        for sec_name, sec_data in proj["sections"].items():
            entry["sections"][sec_name] = sec_data

        # Extract phase display
        bfa_sec = proj["sections"].get("bfa_phase", {})
        bfa_text = bfa_sec.get("text", "")
        if bfa_text:
            phase_line = bfa_text.split("\n")[0]
            phase_line = re.sub(
                r"^(?:Project Status|BFA (?:Project )?Status)\s*:?\s*",
                "", phase_line, flags=re.IGNORECASE,
            ).strip()
            if phase_line:
                entry["bfa_phase_display"] = phase_line

        repo.upsert(entry, allow_curated=True)
        updated += 1
    else:
        # New project
        fields = {"ownership": "curated"}
        try:
            parsed = _parse_header_fields(header_text) or {}
            fields.update(parsed)
        except Exception:
            pass

        fingerprint = (
            f"{fields.get('client', '?')}|"
            f"{fields.get('project_name', '?')}|"
            f"{fields.get('city', '?')}"
        )
        uid = str(uuid.uuid5(uuid.NAMESPACE_DNS, fingerprint))
        slug_raw = (
            f"bfa-{fields.get('client', '')}-"
            f"{fields.get('project_name', '')}-"
            f"{fields.get('city', '')}"
        )
        slug = re.sub(r"[^a-z0-9]+", "-", slug_raw.lower()).strip("-")[:50]
        slug = f"{slug}-{uid[:8]}"

        phase_display = "TBC"
        bfa_sec = proj["sections"].get("bfa_phase", {})
        bfa_text = bfa_sec.get("text", "")
        if bfa_text:
            phase_line = bfa_text.split("\n")[0]
            phase_line = re.sub(
                r"^(?:Project Status|BFA (?:Project )?Status)\s*:?\s*",
                "", phase_line, flags=re.IGNORECASE,
            ).strip()
            if phase_line:
                phase_display = phase_line

        status = "active"
        if "ON HOLD" in header_text.upper():
            status = "on_hold"

        entry = {
            "uid": uid,
            "slug": slug,
            "fingerprint": fingerprint,
            "type": "project",
            "status": status,
            "source": "gdocs",
            "fields": fields,
            "header": {"text": header_text, "html": header_html},
            "sections": proj["sections"],
            "bfa_phase_display": phase_display,
            "bfa_phase_canonical": "TBC",
            "next_steps": [],
        }

        repo.upsert(entry, allow_curated=True)
        created += 1
        print(f"  NEW: {header_text[:80]}")

print(f"\nDone: {updated} updated, {created} created, {skipped} skipped")
print(f"Total: {updated + created} projects + 2 preambles, all ownership=curated")
